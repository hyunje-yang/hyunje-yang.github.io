#!/usr/bin/env python3
"""
build_cv.py -- content/cv.yaml 로 CV docx 와 PDF 를 만든다.

  templates/cv-template.docx  (기존 CV 파일. 서식·스타일·번호매기기의 원본)
        + content/cv.yaml     (내용)
        -> assets/pdf/Hyunje-Yang-CV.docx
        -> assets/pdf/Hyunje-Yang-CV.pdf   (Windows Word 로 변환)

서식은 기존 CV 에서 그대로 가져온다:
  이름 20pt bold 가운데 / 연락처 10pt 가운데
  섹션 제목 11pt bold + 아래 회색 가로선 (Liniapozioma 스타일)
  기관 10.5pt bold + 오른쪽 끝에 지역
  역할 10pt italic + 오른쪽 끝에 기간
  프로젝트 10pt + 점(●) 목록 + 오른쪽 끝에 기간
  본문 10pt + 대시(–) 목록, 양쪽정렬
오른쪽 끝 항목은 원본이 탭을 여러 번 눌러 맞춰 두었는데,
여기서는 오른쪽 탭 스톱 하나로 정확히 정렬한다.
"""
import shutil
import subprocess
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = Path(__file__).resolve().parent
CONTENT = ROOT / "content"
TEMPLATE = ROOT / "templates" / "cv-template.docx"
OUT_DIR = ROOT / "assets" / "pdf"
OUT_DOCX = OUT_DIR / "Hyunje-Yang-CV.docx"
OUT_PDF = OUT_DIR / "Hyunje-Yang-CV.pdf"

RIGHT_EDGE = Inches(7.5)          # 용지 8.5in - 좌우 여백 0.5in
LINK_COLOR = "199CFF"             # 기존 CV 가 쓰던 파란색 (RGB 25,156,255)
TIMES = "/mnt/c/Windows/Fonts/times.ttf"          # 폭 계산용 (Windows 기본 폰트)
TIMES_BOLD = "/mnt/c/Windows/Fonts/timesbd.ttf"
TIMES_ITALIC = "/mnt/c/Windows/Fonts/timesi.ttf"
NUM_DASH = 3                      # 원본 numbering: 대시(–) 목록
NUM_DOT = 6                       # 원본 numbering: 점(●) 목록


# --------------------------------------------------------------------------- 뼈대
def load(name):
    with open(CONTENT / name, encoding="utf-8") as f:
        return yaml.safe_load(f)


def clear_body(doc):
    """템플릿의 스타일·번호매기기·여백은 남기고 본문만 비운다."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_list(p, num_id, ilvl=0):
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    for tag, val in (("w:ilvl", str(ilvl)), ("w:numId", str(num_id))):
        e = OxmlElement(tag)
        e.set(qn("w:val"), val)
        numPr.append(e)
    pPr.append(numPr)


def para(doc, *, size=10, bold=False, italic=False, align="both",
         space_before=0, space_after=0, style=None, num=None,
         left_indent=None, hanging=None, keep_next=False, keep_together=True):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = {"both": WD_ALIGN_PARAGRAPH.JUSTIFY,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        pf.first_line_indent = -hanging
    if num:
        set_list(p, num)
    # 한 항목이 페이지 경계에서 두 쪽으로 갈라지지 않게 한다.
    pf.keep_together = keep_together        # 이 문단의 줄들을 같은 쪽에
    if keep_next:
        pf.keep_with_next = True            # 다음 문단과도 같은 쪽에
    p._default = dict(size=size, bold=bold, italic=italic)
    return p


def run(p, text, *, size=None, bold=None, italic=None, underline=False):
    d = getattr(p, "_default", {"size": 10, "bold": False, "italic": False})
    r = p.add_run(text)
    r.font.size = Pt(size if size is not None else d["size"])
    r.bold = d["bold"] if bold is None else bold
    r.italic = d["italic"] if italic is None else italic
    r.underline = underline
    return r


def link(p, url, text=None, size=10):
    """기존 CV 와 같은 파란 밑줄 하이퍼링크."""
    text = text or url
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), LINK_COLOR); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    for tag in ("w:sz", "w:szCs"):
        e = OxmlElement(tag); e.set(qn("w:val"), str(int(size * 2))); rPr.append(e)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    r.append(t)
    h.append(r)
    p._p.append(h)


INTEREST_COL2 = Inches(3.6)          # RESEARCH INTERESTS 둘째 칸의 점(●) 자리
INTEREST_COL2_TEXT = Inches(3.8)     # 그 글자가 시작하는 자리 (첫째 칸과 같은 0.2in 간격)


def dot_marker(p):
    """줄 가운데에 직접 찍는 ● 표시.
    번호매기기(numId 6)가 쓰는 것과 같은 Wingdings 글자·같은 12pt 라서
    첫째 칸의 점과 크기·모양이 정확히 같다."""
    r = p.add_run("\uF09F")
    rPr = r._element.get_or_add_rPr()
    f = OxmlElement("w:rFonts")
    for tag in ("w:ascii", "w:hAnsi"):
        f.set(qn(tag), "Wingdings")
    rPr.append(f)
    r.font.size = Pt(12)


def right_tab(p):
    """오른쪽 끝에 붙일 항목 앞에 넣는 탭. 위치는 본문 오른쪽 끝."""
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_EDGE, WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")


_FONT_CACHE = {}


def text_width_pt(text, size=10, bold=False, italic=False):
    """Times New Roman 으로 이 글자들이 차지하는 폭(pt)을 실제로 잰다."""
    from PIL import ImageFont
    path = TIMES_BOLD if bold else (TIMES_ITALIC if italic else TIMES)
    key = (path, 200)
    if key not in _FONT_CACHE:
        try:
            _FONT_CACHE[key] = ImageFont.truetype(path, 200)
        except OSError:
            _FONT_CACHE[key] = None
    f = _FONT_CACHE[key]
    if f is None:                       # 폰트를 못 찾으면 글자 수로 어림한다
        return len(text) * size * 0.5
    return f.getlength(text) / 200 * size


SQUEEZE_LIMIT_PT = 0.3      # 글자당 이만큼까지만 좁힌다 (10pt 기준 약 3%)
MIN_GAP_PT = 26             # 왼쪽 글과 오른쪽 날짜 사이에 두려는 간격
TIGHT_GAP_PT = 12           # 자간을 좁혀서라도 한 줄에 넣을 때 허용하는 최소 간격


def squeeze(p, per_char_pt):
    """문단 전체의 자간을 아주 조금 좁힌다.
    한 줄에 살짝 못 들어가는 항목의 날짜가 다음 줄로 밀리는 것을 막는다."""
    val = -int(round(per_char_pt * 20))          # 1/20 pt 단위
    val = max(val, -int(SQUEEZE_LIMIT_PT * 20))
    if val >= 0:
        return
    for r in p.runs:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(val))
        rPr.append(sp)


def _line_width(text, size, bold, per_char_pt):
    """자간을 per_char_pt 만큼 좁혔을 때 이 글자들의 폭(pt)."""
    return text_width_pt(text, size, bold) - per_char_pt * len(text)


def wrap_lines(text, avail, size=10, bold=False, per_char_pt=0.0):
    """Word 가 이 글을 폭 avail 안에서 어떻게 줄바꿈하는지 흉내 낸다.
    끊을 수 있는 자리는 띄어쓰기 뒤와 붙임표(-) 뒤다."""
    stops = [i + 1 for i, c in enumerate(text) if c in " -"] + [len(text)]
    lines, s = [], 0
    while s < len(text):
        best = None
        for e in stops:
            if e <= s:
                continue
            if _line_width(text[s:e].rstrip(), size, bold, per_char_pt) <= avail:
                best = e
            else:
                break
        if best is None:                     # 한 낱말이 통째로 넘칠 때
            best = next(e for e in stops if e > s)
        lines.append(text[s:best].rstrip())
        s = best
        while s < len(text) and text[s] == " ":
            s += 1
    return lines or [""]


def tail_right(p, left_text, right_text, *, indent_in=0.0, size=10,
               bold=False, italic=None):
    """오른쪽 끝 항목(지역, 기간)을 붙인다.

    - 왼쪽 글이 여러 줄로 넘어가면 그 마지막 줄의 오른쪽 끝에 붙인다.
    - 살짝 넘치면 자간을 조금 좁혀 같은 줄에 넣는다.
    - 그래도 안 되면 줄을 바꾸고, 그 문단은 양쪽정렬을 풀어
      글자 사이가 벌어지지 않게 한다."""
    avail = 7.5 * 72 - indent_in * 72                     # 본문 폭(pt)

    def fits(per_char, gap):
        last = wrap_lines(left_text, avail, size, bold, per_char)[-1]
        need = (_line_width(last, size, bold, per_char)
                + text_width_pt(right_text, size, bold, True)
                - per_char * len(right_text) + gap)
        return need <= avail

    per_char = 0.0
    if not fits(0.0, MIN_GAP_PT):
        for step in range(1, int(SQUEEZE_LIMIT_PT * 20) + 1):
            if fits(step / 20, TIGHT_GAP_PT):
                per_char = step / 20
                break
        else:                                             # 좁혀도 안 들어간다
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[-1].add_break()
    if len(wrap_lines(left_text, avail, size, bold, per_char)) > 1:
        # 여러 줄로 넘어간 문단은 양쪽정렬을 풀어 낱말 사이가 벌어지지 않게 한다.
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right_tab(p)
    run(p, right_text, italic=italic)
    if per_char:
        squeeze(p, per_char)


def two_col(doc, left, right, *, size=10, bold=False, italic=False, keep_next=True, **kw):
    """왼쪽 내용 + 오른쪽 끝 내용(지역, 기간)을 한 줄에."""
    p = para(doc, size=size, bold=bold, italic=italic, keep_next=keep_next, **kw)
    run(p, left)
    if right:
        tail_right(p, left, right, size=size, bold=bold)
    return p


def section(doc, title):
    p = doc.add_paragraph(style="Liniapozioma")      # 아래 회색 가로선이 딸린 스타일
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(8)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pf.keep_together = True
    pf.keep_with_next = True
    r = p.add_run(title)
    r.font.size = Pt(11)
    r.bold = True
    return p


def blank(doc, size=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.add_run("").font.size = Pt(size)


def bullet(doc, text, *, dash=True):
    p = para(doc, size=10, num=NUM_DASH if dash else NUM_DOT)
    run(p, text)
    return p


def dash(period):
    return period.replace(" - ", " – ") if period else period


# --------------------------------------------------------------------------- 본문
def build_docx(profile, cv):
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    # ---------- 머리글 ----------
    p = para(doc, size=20, bold=True, align="center", space_after=2)
    run(p, profile["name"].upper())
    p = para(doc, size=10, align="center", space_after=0)
    run(p, f'{profile.get("cv_address", profile["address"])} | {profile["phone"]}')
    p = para(doc, size=10, align="center", space_after=0)
    link(p, f'mailto:{profile["email_school"]}', profile["email_school"])
    if profile.get("site_url"):
        run(p, " | ")
        link(p, profile["site_url"])

    # ---------- EDUCATION ----------
    section(doc, "EDUCATION")
    prev_school = None
    for e in cv["education"]:
        name = e["school"] + (f' ({e["abbr"]})' if e.get("abbr") else "")
        if name != prev_school:                      # 같은 학교면 헤더를 반복하지 않는다
            two_col(doc, name, e["location"], size=10.5, bold=True)
            prev_school = name
        two_col(doc, e["degree"], dash(e["period"]), size=10, italic=True)
        for n in e.get("notes", []):
            p = para(doc, size=10, italic=True, num=NUM_DOT,
                     left_indent=Inches(0.2), hanging=Inches(0.2))
            run(p, n)
        blank(doc)

    # ---------- RESEARCH INTERESTS ----------
    # 홈 화면과 달리 CV 에는 제목 네 개만 싣는다 (설명 문장은 홈 화면 전용).
    # 한 줄에 두 개씩 놓는다. 둘째 칸은 왼쪽 탭 스톱 하나로 세로를 맞춘다.
    if cv.get("research_interests"):
        section(doc, "RESEARCH INTERESTS")
        titles = [it["title"] for it in cv["research_interests"]]
        for i in range(0, len(titles), 2):
            p = para(doc, size=10, num=NUM_DOT, align="left",
                     left_indent=Inches(0.2), hanging=Inches(0.2))
            for stop in (INTEREST_COL2, INTEREST_COL2_TEXT):
                p.paragraph_format.tab_stops.add_tab_stop(
                    stop, WD_TAB_ALIGNMENT.LEFT)
            run(p, titles[i])
            if i + 1 < len(titles):
                p.add_run("\t")
                dot_marker(p)
                p.add_run("\t")
                run(p, titles[i + 1])
        blank(doc)

    # ---------- RESEARCH EXPERIENCE ----------
    section(doc, "RESEARCH EXPERIENCE")
    for org in cv["research_experience"]:
        name = org["org"] + (f' ({org["abbr"]})' if org.get("abbr") else "")
        two_col(doc, name, org["location"], size=10.5, bold=True)
        for r in org.get("roles") or [{"role": org["role"], "period": org["period"]}]:
            label = r["role"]
            if r.get("location"):            # 기관 위치와 다를 때만 역할 줄에 적는다
                label += " \u00b7 " + r["location"]
            two_col(doc, label, dash(r["period"]), size=10, italic=True)
        if org.get("note") and org.get("note_in_cv", True):
            p = para(doc, size=10, italic=True, keep_next=True)
            run(p, org["note"])
        for prj in org["projects"]:
            p = para(doc, size=10, num=NUM_DOT, keep_next=True,
                     left_indent=Inches(0.2), hanging=Inches(0.2))
            run(p, prj["title"])
            tail_right(p, prj["title"], dash(prj["period"]), indent_in=0.2, italic=True)
            for b in prj["bullets"]:
                bullet(doc, b)
            if prj.get("note"):
                p = para(doc, size=10)
                run(p, "※ " + prj["note"])
        blank(doc)

    # ---------- PUBLICATIONS ----------
    section(doc, "PUBLICATIONS")
    if cv.get("under_review"):
        p = para(doc, size=10, bold=True, space_before=2, keep_next=True)
        run(p, "[Under Review]")
        for x in cv["under_review"]:
            p = para(doc, size=10, num=NUM_DASH, align="left")
            write_ref(p, x, kind="preprint")

    if cv.get("in_preparation"):
        p = para(doc, size=10, bold=True, space_before=4, keep_next=True)
        run(p, "[In Preparation]")
        for x in cv["in_preparation"]:
            p = para(doc, size=10, num=NUM_DASH, align="left")
            write_ref(p, x, kind="in_preparation")

    p = para(doc, size=10, bold=True, space_before=4, keep_next=True)
    run(p, "[Journals]")
    for x in cv["publications"]:
        p = para(doc, size=10, num=NUM_DASH, align="left")
        write_ref(p, x, kind="journal")

    p = para(doc, size=10, bold=True, space_before=4, keep_next=True)
    run(p, "[Books]")
    for b in cv["books"]:
        p = para(doc, size=10, num=NUM_DASH, align="left")
        write_authors(p, b["authors"])
        run(p, " ")
        run(p, b["title"] + ".", italic=True)
        run(p, f' {b["publisher"]}, {b["year"]}. (ISBN: {b["isbn"]})')

    # ---------- CONFERENCE PRESENTATIONS ----------
    section(doc, "CONFERENCE PRESENTATIONS")
    for key, label in (("oral", "[Oral Presentations]"), ("poster", "[Poster Presentations]")):
        p = para(doc, size=10, bold=True, space_before=4, keep_next=True)
        run(p, label)
        for c in cv["conferences"][key]:
            p = para(doc, size=10, num=NUM_DASH, align="left")
            write_authors(p, c["authors"])
            run(p, f' “{c["title"]}” {c["venue"]}, {c["location"]}, {c["date"]}.')
            if c.get("note"):
                run(p, f' {c["note"]}.', italic=True)

    # ---------- PEER-REVIEW SERVICE ----------
    section(doc, "PEER-REVIEW SERVICE")
    for j in cv["peer_review"]:
        p = para(doc, size=10, num=NUM_DASH)
        run(p, j)

    # ---------- REGISTERED PATENTS ----------
    section(doc, "REGISTERED PATENTS")
    for pt in cv["patents"]["list"]:
        p = para(doc, size=10, num=NUM_DASH)
        run(p, pt["title"])
        tail_right(p, pt["title"], pt["date"], indent_in=0.25, italic=True)

    # ---------- TECHNICAL SKILLS ----------
    section(doc, "TECHNICAL SKILLS")
    for s in cv["skills"]:
        p = para(doc, size=10, num=NUM_DASH)
        run(p, s["group"] + ": ", bold=True)
        run(p, s["detail"])

    # ---------- AWARDS AND HONORS ----------
    section(doc, "AWARDS AND HONORS")
    for a in cv["awards"]:
        p = para(doc, size=10, num=NUM_DASH)
        run(p, a["name"])
        tail_right(p, a["name"], a["date"], indent_in=0.25, italic=True)

    # ---------- MUSIC ACTIVITIES ----------
    section(doc, "MUSIC ACTIVITIES")
    m = cv["music"]
    head = f'{m["role"]}, Rock Band “{m["band"]}”'
    p = para(doc, size=10, keep_next=True)
    run(p, head)
    tail_right(p, head, dash(m["period"]), size=10, italic=True)
    for s in m["summary"]:
        bullet(doc, s)
    # 음반 목록은 CV 에서만 한 줄로 줄여 쓴다 (홈페이지에는 표지와 함께 따로 있다).
    disc = "; ".join(
        f'{d["title"]} ({d.get("cv_kind", d["kind"].lower())}, {d["year"]})'
        for d in m["discography"])
    bullet(doc, f"Discography: {disc}.")
    for o in m.get("other_activities", []):
        head = f'{o["role"]}, {o["org"]}'
        p = para(doc, size=10, space_before=4, keep_next=True)
        run(p, head)
        tail_right(p, head, dash(o["period"]), size=10, italic=True)
        if o.get("note"):
            bullet(doc, o["note"])

    # ---------- EXTRACURRICULAR ACTIVITIES ----------
    section(doc, "EXTRACURRICULAR ACTIVITIES")
    for g in cv["extracurricular"]:
        p = para(doc, size=10, space_before=2, keep_next=True)
        run(p, g["group"] + ":", bold=True)
        for e in g["entries"]:
            p = para(doc, size=10, num=NUM_DASH)
            run(p, e["text"])
            tail_right(p, e["text"], dash(e["period"]), indent_in=0.25, italic=True)
            for sub in e.get("sub", []):
                p = para(doc, size=10, left_indent=Inches(0.6))
                run(p, sub)

    return doc


def write_authors(p, authors):
    """저자 목록에서 'Yang, H.' 만 굵게."""
    me = "Yang, H."
    parts = authors.split(me)
    for i, seg in enumerate(parts):
        if i:
            run(p, me, bold=True)
        if seg:
            run(p, seg)


def write_ref(p, x, kind):
    write_authors(p, x["authors"])
    run(p, f' “{x["title"]}” ')
    if kind == "in_preparation":          # 투고 전 원고는 저널 이름이 아직 없다
        run(p, "In preparation.", italic=True)
        if x.get("note"):
            run(p, f' ({x["note"]})')
        return
    run(p, x["venue"], italic=True)
    if kind == "journal":
        tail = f', {x["year"]}'
        if x.get("volume"):
            tail += f', {x["volume"]}'
        if x.get("pages"):
            tail += f': p. {x["pages"]}'
        run(p, tail + ".")
    else:
        run(p, f', {x["year"]}.')
    url = f'https://doi.org/{x["doi"]}' if x.get("doi") else x.get("url")
    if url:
        run(p, " ")
        link(p, url)


# --------------------------------------------------------------------------- PDF
def to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Windows 에 설치된 Word 로 PDF 를 만든다. Word 가 없으면 건너뛴다."""
    def win(p: Path) -> str:
        out = subprocess.run(["wslpath", "-w", str(p)], capture_output=True, text=True)
        return out.stdout.strip()

    ps = f'''
$ErrorActionPreference = "Stop"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open("{win(docx_path)}", [ref]$false, [ref]$true)
$doc.SaveAs([ref]"{win(pdf_path)}", [ref]17)
$doc.Close([ref]$false)
$word.Quit()
'''
    r = subprocess.run(["powershell.exe", "-NoProfile", "-Command", ps],
                       capture_output=True, text=True, timeout=300)
    if r.returncode != 0 or not pdf_path.exists():
        print("  PDF 변환 실패:", (r.stderr or r.stdout).strip().splitlines()[:3])
        return False
    return True


def main():
    profile = load("profile.yaml")
    cv = load("cv.yaml")
    if not TEMPLATE.exists():
        print(f"템플릿이 없습니다: {TEMPLATE}")
        return 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_docx(profile, cv)
    doc.save(str(OUT_DOCX))
    print(f"  만듦  assets/pdf/{OUT_DOCX.name}")

    if "--no-pdf" not in sys.argv:
        if to_pdf(OUT_DOCX, OUT_PDF):
            print(f"  만듦  assets/pdf/{OUT_PDF.name}  (Word 변환)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
